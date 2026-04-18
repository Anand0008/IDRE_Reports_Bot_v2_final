"""
Generates IDRE Reports Bot Architecture Document as a .docx file.
Run: python generate_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ───────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Colour palette ─────────────────────────────────────────────────────────────
NAVY    = RGBColor(0x1A, 0x23, 0x4A)   # headings
BLUE    = RGBColor(0x1F, 0x6F, 0xEB)   # accents / epic bars
SILVER  = RGBColor(0xF0, 0xF4, 0xFA)   # table header fill
LIGHT   = RGBColor(0xF7, 0xF9, 0xFF)   # alternating rows
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
DARK    = RGBColor(0x1A, 0x1A, 0x2E)
GREY    = RGBColor(0x6B, 0x72, 0x80)
CODE_BG = RGBColor(0xF3, 0xF4, 0xF6)
CODE_FG = RGBColor(0x1F, 0x6F, 0xEB)

# ── Helper: set cell background ────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, border_type='bottom', color='C7D2FE', sz='6'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    b = OxmlElement(f'w:{border_type}')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), sz)
    b.set(qn('w:space'), '0')
    b.set(qn('w:color'), color)
    tcBorders.append(b)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    tblPr.append(tblBorders)

# ── Style helpers ──────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    # Left rule accent
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '1F6FEB')
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = NAVY
    run.font.name = 'Calibri'
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    run.font.name = 'Calibri'
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    run.font.name = 'Calibri'
    return p

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    # Parse **bold** inline
    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK
        run.font.name = 'Calibri'
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F4F6')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = CODE_FG
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C7D2FE')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(table)
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, NAVY)
        set_cell_border(cell, 'bottom', '1F6FEB', '8')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = LIGHT if r_idx % 2 == 0 else WHITE
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            set_cell_border(cell, 'bottom', 'E2E8F0', '4')
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            parts = str(val).split('**')
            for i, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.bold = (i % 2 == 1)
                run.font.size = Pt(9.5)
                run.font.color.rgb = DARK
                run.font.name = 'Calibri'
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def epic_header(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1A234A')
    p._p.get_or_add_pPr().append(shd)
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run(f"  {text}")
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = WHITE
    run.font.name = 'Calibri'
    return p

def story_header(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EEF2FF')
    p._p.get_or_add_pPr().append(shd)
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run(f"  {text}")
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    run.font.name = 'Calibri'
    return p

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after  = Pt(0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("IDRE Reports Bot")
run.font.name = 'Calibri'
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run("Architecture & Approach Document")
run.font.name = 'Calibri'
run.font.size = Pt(18)
run.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(2)
run = p.add_run("Multi-Agent NL-to-SQL Chatbot  ·  IDRE Dispute Resolution Platform")
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = GREY

doc.add_paragraph().paragraph_format.space_after = Pt(18)

meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(meta_table)
meta = [
    ("System",   "IDRE Reports Bot"),
    ("Platform", "IDRE Dispute Resolution"),
    ("Version",  "1.0 — Phase 1 MVP"),
    ("Date",     "March 2026"),
]
for i, (k, v) in enumerate(meta):
    row = meta_table.rows[i]
    for cell in row.cells:
        set_cell_bg(cell, WHITE)
    kc = row.cells[0]
    vc = row.cells[1]
    kc.width = Inches(1.8)
    vc.width = Inches(3.0)
    rk = kc.paragraphs[0].add_run(k)
    rk.bold = True
    rk.font.size = Pt(10)
    rk.font.color.rgb = NAVY
    rk.font.name = 'Calibri'
    rv = vc.paragraphs[0].add_run(v)
    rv.font.size = Pt(10)
    rv.font.color.rgb = DARK
    rv.font.name = 'Calibri'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Executive Summary")
body(
    "The IDRE Reports Bot is a multi-agent natural language to SQL system that allows platform staff "
    "to query dispute and payment data in plain English, without writing SQL. Users describe what they "
    "want to know — \"how many cases are in initial eligibility review?\" — and the system translates "
    "that into a safe, validated SQL query, executes it against the IDRE MySQL database, and returns a "
    "structured, human-readable response with charts, explanations, and follow-up suggestions."
)
body(
    "The system is built on a LangGraph orchestration layer with specialised agents: Schema Mapper, SQL Writer, "
    "SQL Validator, Executor, and Response Formatter. A Debugger Agent handles retries on failed execution. "
    "All agents share a single GraphState object — no agent receives raw user input directly. Safety is enforced "
    "deterministically at the validator and permission layers, not via LLM judgment."
)
body(
    "The MVP delivers a fully working Streamlit frontend, end-to-end pipeline, and coverage of the 18 core "
    "IDRE reporting metrics using Gemini 2.5 Pro for SQL generation and sentence-transformers (all-MiniLM-L6-v2) "
    "for local schema embeddings. The architecture is designed to evolve incrementally to a production-grade "
    "Next.js frontend, Redis session memory, cross-session intent storage, and column-level access control."
)
divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — CONTEXT
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Context")

h2("Problem")
body(
    "IDRE platform staff — operations managers, financial auditors, eligibility analysts, and executives — "
    "need regular insight into dispute volume, payment status, and case pipeline health. Currently, this "
    "requires either writing manual SQL queries or waiting for engineering to run ad-hoc reports. "
    "Neither approach scales."
)

h2("Database")
add_table(
    ["Property", "Value"],
    [
        ["Host", "AWS RDS MySQL 8 (stage cluster)"],
        ["Database", "idre_stage"],
        ["Scale", "44 tables, ~1.83M rows, 83 FK relationships"],
        ["Central table", "`case` (50 columns, 41K rows) — full dispute lifecycle"],
        ["Largest table", "`email_job` (259K rows, 2.4 GB)"],
        ["SSL", "Required via global-bundle.pem"],
    ],
    col_widths=[2.0, 4.8]
)

h2("Core Reporting Needs")
body(
    "The system must answer 18 verified reporting metrics spanning dispute volume, case status pipeline, "
    "payment tracking, arbitration outcomes, ineligibility counts, and administrative closure. These are "
    "pre-catalogued in metric_cards.json with verified SQL templates and NL triggers."
)

h2("User Personas")
add_table(
    ["Persona", "Code", "Focus"],
    [
        ["Executive Summarizer", "ES", "High-level counts, MTD volume"],
        ["Operations Manager",   "OM", "Overdue and stuck cases"],
        ["Financial Auditor",    "FA", "Payment amounts, unpaid balances"],
        ["Eligibility Analyst",  "CEA","Ineligible and eligibility review cases"],
        ["Data Quality Debugger","DQD","Anomalies, discrepancies"],
    ],
    col_widths=[2.2, 0.9, 3.7]
)
divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — APPROACH
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Approach")
approach_points = [
    "**Safety first, LLM second.** Validation and permission enforcement are fully deterministic. The LLM is used only for schema retrieval (embeddings) and SQL generation — never for safety decisions.",
    "**Pre-verified SQL fast path.** The 18 core metric queries are stored in metric_cards.json with NL trigger phrases. Matching queries bypass the LLM entirely, reducing latency and cost.",
    "**Schema-aware generation.** The SQL Writer receives only the relevant subset of the 44-table schema — retrieved via vector similarity search — rather than the full catalog. This reduces token cost and hallucination surface.",
    "**State flows, not function calls.** All agents operate on a shared GraphState TypedDict. Each agent returns a mutated copy of the state. No agent holds mutable local state between requests.",
    "**Deterministic retry loop.** Execution failures are classified and routed back to the SQL Writer with error context and the failing SQL. Max 2 retries. After that, the pipeline returns a structured failure response.",
    "**Audit everything asynchronously.** Every query event — input, generated SQL, validated SQL, result shape, per-agent latency, retry count — is logged asynchronously so it does not block the response path.",
    "**Access control at two points.** Table permissions are enforced at schema retrieval (ChromaDB filter) and again at SQL validation. Defense in depth.",
    "**Start simple, evolve the stack.** Streamlit for MVP. ChromaDB for local vector storage. In-memory session history. Each has a documented upgrade path (Next.js, pgvector/Pinecone, Redis) that does not require pipeline rewrites.",
]
for pt in approach_points:
    bullet(pt)
divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — CORE COMPONENTS
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Core Components / Modules")

components = [
    ("4.1  State Layer — state/context.py",
     "Define the single data contract that flows through the entire pipeline.",
     [
         "GraphState TypedDict with all fields: user query, session ID, relevant tables, schema context, generated SQL, validated SQL, query result, row count, execution error, formatted response, retry count, error message, agent trace.",
         "All agents read from and write to this object. No side channels.",
         "Agent trace field accumulates a step-by-step log of what each agent did within a single request.",
     ]),
    ("4.2  Configuration — config/",
     "Centralised environment and metadata management.",
     [
         "settings.py — Pydantic-Settings loader for all env vars (DB credentials, SSL path, Gemini API key — Gemini_API_Key).",
         "schema_catalog.json — Machine-generated full schema: 44 tables, columns, types, FKs, sample values. Source of truth for the validator and schema mapper.",
         "metric_cards.json — 18 pre-verified metric queries with NL triggers, SQL templates, and persona tags.",
         "permissions.json — Role-to-table mapping for access control enforcement.",
     ]),
    ("4.3  Database Connector — db/connector.py",
     "Manage the SQLAlchemy engine lifecycle.",
     [
         "Singleton engine instantiated once at startup.",
         "SSL-enforced connection to AWS RDS MySQL 8 using global-bundle.pem.",
         "Connection pool with pre-ping and 300-second recycle.",
         "All queries run via the read-only DB user app_idre_rw — no write access at the infrastructure level.",
     ]),
    ("4.4  Schema Mapper Agent — agents/schema_mapper.py",
     "Translate user intent into a relevant schema subset.",
     [
         "Embeds 44 table documents (name + columns + FKs + sample values) into ChromaDB using sentence-transformers all-MiniLM-L6-v2 (runs locally, no API key required).",
         "Persists the vector index to data/chroma_db/ — built once, reused across restarts.",
         "On each query, runs top-K similarity search (default K=6) filtered to the user's permitted table set.",
         "Returns full schema context string (columns, types, FKs, sample values) for the retrieved tables.",
     ]),
    ("4.5  SQL Writer Agent — agents/sql_writer.py",
     "Produce a correct SELECT query for the user's intent.",
     [
         "Fast path: fuzzy-match user query against nl_triggers in metric_cards.json. If matched, return pre-verified SQL without calling the LLM.",
         "LLM path: call Gemini 2.5 Pro at temperature 0 with a structured system prompt containing schema context, rules, and error context on retries.",
         "On retry, inject the previous failed SQL and the classified execution error into the prompt.",
         "Strip markdown fences from LLM output defensively.",
     ]),
    ("4.6  SQL Validator Agent — agents/sql_validator.py",
     "Deterministic pre-execution safety gate — no LLM involved.",
     [
         "Block any query not starting with SELECT.",
         "Regex-block DDL/DML keywords: INSERT, UPDATE, DELETE, DROP, CREATE, ALTER, TRUNCATE, EXEC, CALL, GRANT, REVOKE.",
         "Block multi-statement queries (semicolon detection). Detect comment-injection and UNION-based patterns.",
         "Extract all table references from FROM and JOIN clauses; verify each exists in schema_catalog.json and the user's permitted set.",
     ]),
    ("4.7  Executor — agents/executor.py",
     "Run validated SQL safely and return structured results.",
     [
         "Enforce LIMIT 1000 on any non-aggregate query missing a LIMIT clause.",
         "Set MAX_EXECUTION_TIME session variable to 30 seconds per query.",
         "Execute via SQLAlchemy; return results as List[Dict[str, Any]].",
         "Coerce non-JSON-serialisable types (datetime, Decimal) at fetch time. Return execution errors as structured strings.",
     ]),
    ("4.8  Debugger Agent — core/orchestrator.py",
     "Classify errors and direct structured retries.",
     [
         "Classify execution errors into: INVALID_COLUMN, UNKNOWN_TABLE, SYNTAX_ERROR, TIMEOUT, PERMISSION_DENIED, EMPTY_RESULT.",
         "Route back to SQL Writer with error class, raw error message, and failing SQL injected into the retry prompt.",
         "Increment retry_count in GraphState on each retry pass. Circuit break after 2 retries.",
     ]),
    ("4.9  Response Formatter — agents/response_formatter.py",
     "Convert raw query results into a human-readable, structured response.",
     [
         "Auto-select output format: single bold number (single aggregate), markdown table (2–50 rows), truncated table with row count notice (50+ rows).",
         "Flag time-series and grouped-count results for chart rendering.",
         "Append the SQL used, collapsibly, to every response. Generate a one-sentence Query Explainer.",
         "Generate 2–3 proactive follow-up suggestion chips.",
     ]),
    ("4.10  Orchestrator — core/orchestrator.py",
     "Assemble and run the LangGraph pipeline.",
     [
         "Declare all agents as LangGraph nodes with typed edges.",
         "Conditional routing: validator fail → END; executor fail + retries remaining → increment_retry → sql_writer; max retries → END.",
         "Compile the pipeline once at module import as a singleton.",
         "Expose run_query(query, session_id) as the single entry point.",
     ]),
    ("4.11  Frontend — app.py",
     "Streamlit chat interface for MVP.",
     [
         "Persistent chat message history per session.",
         "Sidebar: API key input, SQL display toggle, agent trace toggle, session reset.",
         "Render markdown tables, single-value callouts, charts, and follow-up suggestion chips.",
     ]),
]
for title, role, bullets in components:
    h3(title)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run("Role: ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = NAVY
    r1.font.name = 'Calibri'
    r2 = p.add_run(role)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK
    r2.font.name = 'Calibri'
    for b in bullets:
        bullet(b)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — LOGIC APPROACH
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Logic Approach")
body(
    "The system's reasoning is structured as a sequential enrichment pipeline, not a free-form agent loop. "
    "Each stage narrows ambiguity and adds structure until a safe, executable query is produced."
)
stages = [
    ("Stage 1 — Ambiguity Resolution",
     "Before any schema lookup occurs, the Clarification Agent scores the query for entity, temporal, metric, and scope ambiguity. High-ambiguity queries pause the pipeline and return a question to the user. Moderate-ambiguity queries proceed with an assumption annotation appended to the final response."),
    ("Stage 2 — Intent to Schema",
     "The Schema Mapper converts the (clarified) user intent into a filtered schema context — a subset of the 44 tables most semantically relevant to the query, limited to the user's permitted set. Business glossary terms are resolved before this step, so domain language is grounded into SQL-ready expressions."),
    ("Stage 3 — Schema to SQL",
     "The SQL Writer receives structured schema context, not the raw user query against a full schema. It first checks metric cards for a pre-verified match. On a miss, it calls Gemini 2.5 Pro with schema context, rules, and (on retries) prior error context. The LLM is working within a tightly constrained input space."),
    ("Stage 4 — Deterministic Safety",
     "The Validator runs a fully deterministic gauntlet: keyword blocking, table existence, column existence, permission re-check, multi-statement detection. If any check fails, the pipeline ends or retries — the LLM is never consulted for safety decisions."),
    ("Stage 5 — Safe Execution",
     "The Executor runs the validated SQL with hard resource limits. Results are serialised and returned as typed dicts. Errors are classified and routed to the retry loop."),
    ("Stage 6 — Structured Response",
     "The Response Formatter selects the appropriate output format based on result shape, generates an explainer, and surfaces follow-up suggestions. The user receives a complete response: result + explanation + SQL + suggested next queries."),
]
for title, desc in stages:
    h3(title)
    body(desc)
divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
h1("6. System Architecture")

h2("Pipeline Flow")
code_block(
"""User (Streamlit UI)
         │
         ▼
  [app.py — Streamlit]  ──  Session state (in-memory per session)
         │
         ▼
  [Context Loader]  ·  Build GraphState  ·  Attach session_id
         │
         ▼
  [1. Schema Mapper]  ·  ChromaDB vector search (sentence-transformers)
         │
         ▼
  [2. SQL Writer]  ·  metric_cards fast path  ·  Gemini 2.5 Pro fallback
         │
         ▼
  [3. SQL Validator]  ·  DDL/DML block  ·  Table check  ·  Permission re-check
         │
         ├── VALID ──►  [4. Executor]  ·  LIMIT  ·  Timeout  ·  Serialise
         │                    │
         │              ┌─────┴───────────────────┐
         │              │ ERROR (retry_count < 2)  │
         │              │  Classify → SQL Writer   │
         │              └─────────────────────────┘
         │              SUCCESS
         │                    │
         │                    ▼
         │             [5. Response Formatter]
         │
         ▼
  [Response → Streamlit chat]"""
)

h2("Component Technology Map")
add_table(
    ["Layer", "Component", "Technology"],
    [
        ["Orchestration",  "LangGraph StateGraph",      "langgraph 1.1.3"],
        ["LLM",            "SQL Writer",                 "Gemini 2.5 Pro via langchain-google-genai"],
        ["Embeddings",     "Schema Mapper",              "all-MiniLM-L6-v2 (sentence-transformers, local)"],
        ["Vector Store",   "Schema index",               "ChromaDB (dev) → pgvector (prod)"],
        ["Database",       "Query execution",            "SQLAlchemy + mysql-connector-python"],
        ["Backend",        "API layer",                  "Streamlit (MVP) — FastAPI planned"],
        ["Session",        "Within-session memory",      "In-memory dict (current) → Redis (planned)"],
        ["Frontend",       "Chat UI",                    "Streamlit (MVP) → Next.js (planned)"],
        ["Config",         "Schema + permissions",       "JSON files"],
        ["Audit",          "Query logging",              "Async DB write / structured log (planned)"],
    ],
    col_widths=[1.6, 2.2, 2.9]
)
divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — RUNTIME FLOW
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Runtime Flow — Single Request")
body('Input: "How many cases are in initial eligibility review?"  |  User role: Operations Manager')
steps = [
    ("Step 1 — Context Loader",
     "Initialise GraphState with user_query and session_id. Build initial state object passed to the pipeline."),
    ("Step 2 — Schema Mapper",
     "Embed query using all-MiniLM-L6-v2 (sentence-transformers, local). Top-6 similar tables retrieved from ChromaDB: case, case_action, case_party, case_note, organization, user. Build schema context string with columns, FKs, and sample status values for the case table."),
    ("Step 3 — SQL Writer (Fast Path)",
     'Match "initial eligibility" against nl_triggers in metric card #4. Return pre-verified SQL: SELECT COUNT(*) AS initial_eligibility_review FROM `case` WHERE status = \'INITIAL_ELIGIBILITY_REVIEW\'. No LLM call.'),
    ("Step 4 — SQL Validator",
     "Starts with SELECT ✓  |  No DDL/DML keywords ✓  |  Single statement ✓  |  Table `case` exists ✓  →  Validated. Pass to Executor."),
    ("Step 5 — Executor",
     "Query is aggregate (COUNT, no GROUP BY) — no LIMIT appended. Set MAX_EXECUTION_TIME = 30000. Execute. Returns: [{\"initial_eligibility_review\": 847}] in 210ms."),
    ("Step 6 — Response Formatter",
     "Single aggregate value detected. Output: initial_eligibility_review: 847. SQL block appended to response (shown/hidden by user toggle in Streamlit sidebar)."),
    ("Step 7 — Response to Frontend",
     "Streamlit renders the result. Agent trace available in sidebar toggle for debugging. Session message history updated."),
]
for title, desc in steps:
    h3(title)
    body(desc)
divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — SUPPORTING SYSTEMS
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Supporting Systems")

h2("Session Memory")
add_table(
    ["Tier", "Scope", "Storage", "Status"],
    [
        ["Within-session", "Current conversation",  "st.session_state (Streamlit in-memory)", "Implemented — chat history per session"],
        ["Cross-session",  "Saved queries",          "DB table (planned)",                    "Planned — not yet implemented"],
        ["Intent memory",  "Common patterns",        "DB table (planned)",                    "Planned — not yet implemented"],
    ],
    col_widths=[1.4, 1.8, 2.2, 1.9]
)

h2("Access Control")
body(
    "Three enforcement points provide defense in depth:"
)
bullet("**Schema Retrieval** — ChromaDB query filtered to permitted table set. Prevents schema context leakage to SQL Writer.")
bullet("**SQL Validation** — Tables and columns in generated SQL verified against permitted set. Blocks attempts to query unpermitted tables even via prompt manipulation.")
bullet("**Execution** — Read-only DB user at the infrastructure level. Even if both upstream checks fail, the DB user cannot modify data.")
body("Column-level restrictions (e.g. bank_account.accountNumber, user.email) are enforced by stripping sensitive columns from schema context and by explicit column-level check in the Validator.")

h2("Business Glossary")
add_table(
    ["Term", "SQL Expression"],
    [
        ['"unpaid"',    "payment.status = 'PENDING'"],
        ['"MTD"',       "createdAt >= DATE_FORMAT(CURDATE(), '%Y-%m-01')"],
        ['"default"',   "status IN ('CLOSED_DEFAULT_IP', 'CLOSED_DEFAULT_NIP')"],
        ['"overdue"',   "dueDateAt < NOW() AND status NOT IN (...)"],
    ],
    col_widths=[1.8, 5.0]
)

h2("Audit Trail")
body("Every query event writes asynchronously to query_audit_log:")
code_block(
    "session_id | user_id | user_role | raw_query | generated_sql | validated_sql\n"
    "row_count  | total_latency_ms | per_agent_latency | retry_count\n"
    "fast_path_hit | error_class | final_status | timestamp"
)
body("Security audit events (blocked DDL, permission denials) write to a separate security_audit_log and can trigger Slack webhook alerts on threshold breach.")
divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — SUMMARY TABLE
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Summary Table — Key Decisions")
add_table(
    ["Decision", "Choice", "Rationale"],
    [
        ["Orchestration",       "LangGraph",                "Native graph-with-cycles for retry loop; typed state; explicit edge control"],
        ["LLM for SQL",         "Gemini 2.5 Pro",           "High accuracy for complex SQL; temperature 0; constrained schema context"],
        ["Embeddings",          "all-MiniLM-L6-v2 (local)", "No API key needed; runs locally via sentence-transformers; persisted to ChromaDB"],
        ["Vector Store",        "ChromaDB (dev)",           "Zero-infrastructure; swap to pgvector/Pinecone without pipeline changes"],
        ["Validation",          "Fully deterministic",      "LLM judgment is not acceptable for security-critical safety checks"],
        ["Permission enforcement","Two-point (retrieval + validation)", "Schema leakage prevention + final SQL check"],
        ["Fast path",           "metric_cards.json",        "18 core metrics cover majority of daily queries; zero LLM cost"],
        ["Session storage",     "In-memory (Streamlit)",    "Simple for MVP; Redis upgrade path planned — no pipeline changes required"],
        ["Frontend",            "Streamlit (MVP)",          "Rapid iteration; Next.js + Tailwind documented as upgrade path"],
        ["DB connection",       "Read-only MySQL user",     "Infrastructure-level safety; cannot be bypassed by app failures"],
        ["Row cap",             "LIMIT 1000 + 30s timeout", "Prevents runaway queries on large tables (email_job: 259K rows)"],
        ["Retry strategy",      "Max 2, error-context injection", "Covers fixable SQL errors; bounded cost; structured failure after cap"],
    ],
    col_widths=[2.0, 1.8, 3.0]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — CURRENT IMPLEMENTATION STATUS
# ══════════════════════════════════════════════════════════════════════════════
h1("10. Current Implementation Status (Phase 1 MVP)")
body(
    "The following components have been built and verified end-to-end as of March 2026. "
    "All items below are working in the staging environment against idre_stage (AWS RDS MySQL 8, 44 tables, ~1.83M rows)."
)

h2("Implemented and Working")
add_table(
    ["Component", "File(s)", "Notes"],
    [
        ["GraphState TypedDict",         "state/context.py",               "All fields defined; flows through full pipeline"],
        ["Pydantic-Settings config",     "config/settings.py",             "Loads .env; case-insensitive key matching; Gemini_API_Key supported"],
        ["SQLAlchemy DB connector",      "db/connector.py",                "Singleton engine; SSL; read-only user app_idre_rw; connection pool"],
        ["Schema Mapper Agent",          "agents/schema_mapper.py",        "sentence-transformers all-MiniLM-L6-v2; ChromaDB collection schema_tables_v2; top-6 retrieval"],
        ["SQL Writer Agent",             "agents/sql_writer.py",           "metric_cards fast path + Gemini 2.5 Pro fallback; retry-aware"],
        ["SQL Validator Agent",          "agents/sql_validator.py",        "DDL/DML block; table existence check; multi-statement guard"],
        ["Executor Agent",               "agents/executor.py",             "LIMIT 1000; 30s timeout; datetime/Decimal coercion"],
        ["Response Formatter",           "agents/response_formatter.py",   "Single value / table / truncated; SQL block appended"],
        ["LangGraph Orchestrator",       "core/orchestrator.py",           "Full pipeline; retry loop max 2; run_query() entry point"],
        ["Streamlit Frontend",           "app.py",                         "Chat UI; SQL toggle; agent trace toggle; session reset"],
        ["Metric Cards (18 queries)",    "config/metric_cards.json",       "18 pre-verified SQLs with NL triggers; verified against email reports"],
        ["Schema Catalog",               "config/schema_catalog.json",     "44 tables; columns, types, FKs, sample values"],
        ["Audit Findings Report",        "IDRE_Report_Audit_Findings.md",  "7 email dates compared; Gemini exact match 8/17; 9/17 need metric cards"],
    ],
    col_widths=[2.1, 2.1, 2.6]
)

h2("Planned — Not Yet Implemented")
add_table(
    ["Feature", "Epic / Story", "Notes"],
    [
        ["Clarification Agent (ambiguity scoring)",    "Epic 3",     "Pipeline proceeds directly to Schema Mapper; no ambiguity check yet"],
        ["Cross-session saved queries",                "Story 2.3",  "Session history is in-memory only; not persisted across restarts"],
        ["Business glossary term resolution",          "Story 4.2",  "Terms like 'unpaid', 'MTD' not pre-resolved; handled by LLM in prompt"],
        ["Few-shot example library",                   "Story 5.3",  "Not yet injected into SQL Writer prompt"],
        ["Chart rendering",                            "Story 8.2",  "Markdown tables only; no chart output yet"],
        ["Query Explainer & follow-up chips",          "Stories 8.3, 8.4", "Not yet generated by Response Formatter"],
        ["Async audit trail logging",                  "Epic 11",    "No query_audit_log DB table; agent_trace written to session state only"],
        ["FastAPI backend",                            "Story 12.2", "Streamlit calls orchestrator directly; no REST layer yet"],
        ["Next.js frontend",                           "Phase 2",    "Streamlit MVP only"],
        ["Redis session store",                        "Story 2.2",  "Streamlit st.session_state in-memory only"],
        ["Column-level access control",                "Story 10.3", "Table-level validation only; column restrictions not yet enforced"],
        ["Result caching",                             "Story 7.4",  "No caching; every query hits the DB"],
    ],
    col_widths=[2.4, 1.4, 3.0]
)

divider()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — JIRA EPICS & STORIES
# ══════════════════════════════════════════════════════════════════════════════
h1("11. Jira Epics & Stories")

epics = [
    ("Epic 1: Schema Catalog & Embedding Pipeline", [
        ("Story 1.1 – Schema Catalog Build & Maintenance", [
            "Parse all 44 tables: extract column names, types, nullability, PKs, FKs, and sample values into schema_catalog.json.",
            "Store human-written table descriptions alongside auto-extracted metadata.",
            "Script for incremental refresh when schema changes (re-embed only changed tables).",
        ]),
        ("Story 1.2 – Table Embedding & Vector Store Population", [
            "Generate rich text documents per table (name + columns + FK relationships + sample values).",
            "Embed using sentence-transformers all-MiniLM-L6-v2 (runs locally) and persist to ChromaDB (dev) / pgvector (prod).",
            "Build upsert logic so re-runs don't duplicate entries.",
        ]),
        ("Story 1.3 – Join Graph Construction", [
            "Build a static FK-based join graph from schema_catalog.json.",
            "Expose get_join_path(table_a, table_b) utility for the Schema Mapper.",
            "Validate graph completeness: flag tables with no FK connections.",
        ]),
    ]),
    ("Epic 2: Context & Session Management", [
        ("Story 2.1 – StateContext Object Design", [
            "Define GraphState TypedDict with all fields: query, tables, SQL, results, errors, retry count, trace.",
            "Enforce that no agent receives raw user input — all input arrives via StateContext.",
            "Add agent_trace list for step-by-step audit within a single request.",
        ]),
        ("Story 2.2 – Within-Session Memory", [
            "Store last N queries per session_id (in-memory dict for MVP, Redis for production).",
            "Resolve pronouns and references (\"those cases\", \"the same filter\") using session history.",
            "Inject relevant prior query context into the SQL Writer prompt on follow-up queries.",
        ]),
        ("Story 2.3 – Cross-Session Saved Queries", [
            "Persist named queries to DB (query text, SQL, result shape, user, timestamp).",
            "Expose bookmark endpoint: POST /queries/save.",
            "Allow retrieval by name or tag: \"show me my saved MTD query\".",
        ]),
    ]),
    ("Epic 3: Clarification & Ambiguity Handling", [
        ("Story 3.1 – Ambiguity Scorer", [
            "Score each query on four dimensions: entity, temporal, metric, and scope ambiguity.",
            "Return a composite confidence score (0–1).",
            "Threshold configuration: below 0.5 triggers clarification, 0.5–0.75 adds an assumption note.",
        ]),
        ("Story 3.2 – Clarification Agent", [
            "Generate one targeted clarifying question per ambiguous dimension.",
            "Pause the pipeline and return the question with a structured needs_clarification: true flag.",
            "On user response, re-enter pipeline with merged clarification context.",
        ]),
        ("Story 3.3 – Assumption Annotation", [
            "For moderate-ambiguity queries, append \"I assumed X — let me know if you meant something else\".",
            "Log assumed interpretations in the audit trail.",
        ]),
    ]),
    ("Epic 4: Schema Mapping & Retrieval", [
        ("Story 4.1 – Semantic Table Retrieval", [
            "Run vector similarity search over embedded table documents for each incoming query.",
            "Return top-K tables (configurable, default 6) filtered by user's permitted table set.",
            "Cache retrieval results per session for follow-up queries with the same table set.",
        ]),
        ("Story 4.2 – Business Glossary Term Resolution", [
            "Pre-pass over user query to detect glossary terms (e.g. \"unpaid\", \"MTD\", \"default\").",
            "Map each term to its SQL expression or filter clause before schema retrieval.",
            "Inject resolved terms into schema context passed to SQL Writer.",
        ]),
        ("Story 4.3 – Schema Context Builder", [
            "Construct full schema context string for retrieved tables: columns, types, FKs, sample values.",
            "Trim context to fit LLM token budget (prioritise PK/FK columns and high-cardinality enums).",
            "Add join path hints between retrieved tables.",
        ]),
    ]),
    ("Epic 5: SQL Generation", [
        ("Story 5.1 – Metric Card Fast Path", [
            "Match user query against nl_triggers in metric_cards.json using fuzzy string match.",
            "Return pre-verified SQL directly, bypassing LLM entirely.",
            "Track fast-path hit rate in audit logs for coverage monitoring.",
        ]),
        ("Story 5.2 – LLM SQL Writer", [
            "System prompt: schema context + rules (SELECT only, backtick table names, LIMIT 1000, readable aliases).",
            "Model: Gemini 2.5 Pro at temperature 0 (via langchain-google-genai); structured to return only raw SQL.",
            "Strip markdown fences from output as a defensive post-processing step.",
        ]),
        ("Story 5.3 – Few-Shot Example Library", [
            "Manually craft 15–20 representative (query → SQL) pairs covering single-table, join, aggregate, date-filter, status-filter patterns.",
            "Inject dynamically: select 3–5 most similar examples using embedding similarity.",
            "Store in config/few_shot_examples.json for easy iteration.",
        ]),
        ("Story 5.4 – Retry-Aware SQL Rewriting", [
            "On retry, inject execution error and failed SQL into Writer prompt.",
            "Add explicit instruction: \"The previous query failed with [error]. Fix the following query: [sql]\".",
            "Limit rewrite to the specific clause causing the error where possible.",
        ]),
    ]),
    ("Epic 6: SQL Validation & Safety", [
        ("Story 6.1 – DDL/DML Blocker", [
            "Regex-based detection of INSERT, UPDATE, DELETE, DROP, CREATE, ALTER, TRUNCATE, EXEC, CALL, GRANT, REVOKE.",
            "Block and return structured error before any DB connection is made.",
            "Log all blocked attempts to the security audit trail.",
        ]),
        ("Story 6.2 – Table & Column Existence Check", [
            "Extract all table and column names referenced in generated SQL.",
            "Validate each against schema_catalog.json.",
            "Return a structured error listing unknown identifiers for the Debugger to use.",
        ]),
        ("Story 6.3 – Multi-Statement & Injection Guard", [
            "Block queries containing more than one statement (semicolon split check).",
            "Detect and block comment-injection patterns (--, /* */, UNION SELECT).",
            "Normalise whitespace and strip inline comments before all other checks.",
        ]),
        ("Story 6.4 – Permission Re-Check at Validation", [
            "Re-verify that all tables referenced in the final SQL are within the user's permitted set.",
            "Second enforcement point (first is in schema retrieval) — defense in depth.",
            "Return 403 Forbidden with table name if permission is violated.",
        ]),
    ]),
    ("Epic 7: Query Execution", [
        ("Story 7.1 – Read-Only Execution Engine", [
            "Connect via SQLAlchemy using the read-only DB user (app_idre_rw).",
            "SSL enforced via global-bundle.pem; connection pool with pre-ping and 5-minute recycle.",
            "All queries run in a read-only transaction context.",
        ]),
        ("Story 7.2 – Row Cap & Timeout Enforcement", [
            "Append LIMIT 1000 to any non-aggregate query missing a LIMIT clause.",
            "Set MAX_EXECUTION_TIME session variable to 30 seconds per query.",
            "Raise structured QueryTimeoutError on breach.",
        ]),
        ("Story 7.3 – Result Serialisation", [
            "Return results as List[Dict[str, Any]] with column names preserved.",
            "Coerce non-JSON-serialisable types (datetime, Decimal) to string/float at fetch time.",
            "Include metadata: row count, column names, execution time ms.",
        ]),
        ("Story 7.4 – Query Result Caching", [
            "Cache validated SQL → result set for identical queries within a configurable TTL (default 5 min).",
            "Cache key: SHA-256 of normalised SQL + user permission set.",
            "Skip cache on queries containing NOW(), CURDATE(), or RAND().",
        ]),
    ]),
    ("Epic 8: Response Formatting & Explainability", [
        ("Story 8.1 – Output Format Auto-Selection", [
            "Single aggregate value → bold number with label.",
            "2–50 rows → GitHub-flavoured markdown table.",
            "50+ rows → summary (first 50 rows + row count notice). Time-series → chart flag.",
        ]),
        ("Story 8.2 – Chart Rendering", [
            "Rule-based chart type selection: bar (categorical counts), line (time series), pie (proportions <6 categories).",
            "Render using Streamlit native chart components for MVP; Recharts for Next.js Phase 2.",
            "Pass chart config (x-axis, y-axis, title) as structured metadata alongside raw data.",
        ]),
        ("Story 8.3 – Query Explainer", [
            "Generate a one-sentence plain-English description of what the SQL query does.",
            "Include filter conditions, join logic, and aggregation in the explanation.",
            "Display collapsibly beneath the result — not inline — to reduce noise.",
        ]),
        ("Story 8.4 – Proactive Follow-Up Suggestions", [
            "Generate 2–3 natural follow-up question suggestions based on query type and result shape.",
            "Examples: \"Break down by organisation?\" / \"Filter to a specific date range?\"",
            "Render as clickable chips in the UI that populate the input box.",
        ]),
    ]),
    ("Epic 9: Debugging & Retry System", [
        ("Story 9.1 – Error Classification Taxonomy", [
            "Classify execution errors into: INVALID_COLUMN, UNKNOWN_TABLE, SYNTAX_ERROR, TIMEOUT, PERMISSION_DENIED, EMPTY_RESULT.",
            "Each class maps to a specific retry strategy and error message template.",
            "Log class + raw error in audit trail.",
        ]),
        ("Story 9.2 – Debugger Agent", [
            "Receive failed SQL + classified error + schema context.",
            "Construct targeted repair prompt: include error, failing clause, and relevant schema excerpt.",
            "Route back to SQL Writer with repair context; do not re-run Schema Mapper unless table error.",
        ]),
        ("Story 9.3 – Retry Loop & Circuit Breaker", [
            "Max 2 retries; each increments retry_count in StateContext.",
            "After max retries, return structured failure response with the last error and SQL attempted.",
            "Circuit breaker: if same error class repeats across 3 consecutive queries, surface a warning.",
        ]),
    ]),
    ("Epic 10: Access Control", [
        ("Story 10.1 – Role & Permission Model", [
            "Define roles: Admin, Executive, Operations, Finance, Analyst.",
            "Map each role to a permitted table set stored in config/permissions.json.",
            "Roles assigned per user in the session context; no dynamic elevation.",
        ]),
        ("Story 10.2 – Schema Retrieval Filtering", [
            "Filter ChromaDB retrieval results to only return tables within the user's permitted set.",
            "First enforcement point — prevents schema context leakage to the SQL Writer.",
        ]),
        ("Story 10.3 – Column-Level Access Control", [
            "Define sensitive columns per table (e.g. bank_account.accountNumber, user.email).",
            "Strip sensitive columns from schema context before passing to SQL Writer.",
            "Validator rejects queries that reference restricted columns regardless of role.",
        ]),
        ("Story 10.4 – Permission Audit Logging", [
            "Log every permission check: user, role, table requested, outcome, timestamp.",
            "Separate from query audit log — feeds security dashboard.",
            "Alert on repeated denied attempts from the same user within a session.",
        ]),
    ]),
    ("Epic 11: Audit Trail & Observability", [
        ("Story 11.1 – Per-Query Audit Log", [
            "Log: session_id, user_id, raw query, generated SQL, validated SQL, result row count, total latency, per-agent latency, retry count, final status.",
            "Write asynchronously (non-blocking) after response is returned to user.",
            "Store in query_audit_log DB table or append to structured JSON log file for MVP.",
        ]),
        ("Story 11.2 – Agent-Level Latency Tracking", [
            "Record entry/exit timestamps for each agent in the pipeline.",
            "Surface per-agent latency in the agent trace visible in UI debug mode.",
            "Aggregate to identify bottleneck agents in observability dashboard.",
        ]),
        ("Story 11.3 – Error & Anomaly Alerting", [
            "Track error rate per query type and per agent over rolling 1-hour windows.",
            "Alert (log + optional Slack webhook) when error rate exceeds 20% or latency exceeds 15s p95.",
            "Track metric card hit rate vs LLM fallback rate to monitor coverage drift.",
        ]),
        ("Story 11.4 – Usage Analytics Dashboard", [
            "Most queried tables, most common query intents, most triggered metric cards.",
            "Daily active query volume, average retries per query, top failing query patterns.",
            "Built on top of audit log table; Streamlit admin page for MVP.",
        ]),
    ]),
    ("Epic 12: Multi-Agent Orchestration & Integration", [
        ("Story 12.1 – LangGraph Pipeline Assembly", [
            "Wire all agents as LangGraph nodes with explicit typed edges.",
            "Conditional routing: validator fail → END; executor fail → retry → sql_writer; max retries → END.",
            "Compile pipeline once at startup; expose run_query(query, session_id) as the single entry point.",
        ]),
        ("Story 12.2 – FastAPI Backend", [
            "POST /query — accepts user query + session_id, returns formatted response + metadata.",
            "GET /sessions/{session_id}/history — returns query history for the session.",
            "GET /health — pipeline readiness check (DB, vector store, LLM reachability).",
        ]),
        ("Story 12.3 – Streamlit Frontend (MVP)", [
            "Chat interface with persistent message history per session.",
            "Sidebar: API key input, SQL toggle, agent trace toggle, session reset.",
            "Render markdown tables, charts, single-value callouts, and follow-up suggestion chips.",
        ]),
        ("Story 12.4 – End-to-End Integration Testing", [
            "50-query golden test set covering all metric cards, join queries, date filters, edge cases, and adversarial inputs.",
            "Automated test runner: compare generated SQL against expected SQL templates; flag semantic deviations.",
            "Security test suite: SQL injection attempts, cross-tenant access, DDL injection — all must be blocked by validator.",
        ]),
    ]),
]

for epic_title, stories in epics:
    epic_header(epic_title)
    for story_title, story_bullets in stories:
        story_header(story_title)
        for b in story_bullets:
            bullet(b, level=1)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
output_path = "<HOME>/Downloads/IDRE_ReportsBot_Architecture_v2.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
